"""
文档处理模块
支持 txt, pdf, docx, doc, xlsx 等格式
"""

import os
import subprocess
import tempfile
from typing import List, Optional
from pathlib import Path

from langchain_core.documents import Document
import pypdf
from docx import Document as DocxDocument
import pandas as pd
import openpyxl


class DocumentProcessor:
    """文档处理器"""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.docx', '.doc', '.xlsx', '.xls']
    
    def __init__(self, input_dir: str):
        self.input_dir = Path(input_dir)
        self.input_dir.mkdir(parents=True, exist_ok=True)
    
    def load_all_documents(self) -> List[Document]:
        """加载所有文档"""
        documents = []
        
        for file_path in self._find_documents():
            try:
                docs = self.load_document(file_path)
                documents.extend(docs)
                print(f"已加载: {file_path.name}")
            except Exception as e:
                print(f"加载失败 {file_path.name}: {e}")
        
        return documents
    
    def _find_documents(self) -> List[Path]:
        """查找所有支持的文档"""
        files = []
        for ext in self.SUPPORTED_EXTENSIONS:
            files.extend(self.input_dir.glob(f"*{ext}"))
        return files
    
    def load_document(self, file_path: Path) -> List[Document]:
        """加载单个文档"""
        suffix = file_path.suffix.lower()

        loaders = {
            '.txt': self._load_txt,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_doc,
            '.xlsx': self._load_xlsx,
            '.xls': self._load_xlsx,
        }

        loader = loaders.get(suffix)
        if not loader:
            raise ValueError(f"不支持的文件格式: {suffix}")

        content = loader(file_path)

        return [Document(
            page_content=content,
            metadata={
                'source': str(file_path),
                'filename': file_path.name,
                'type': suffix[1:]
            }
        )]
    
    def _load_txt(self, file_path: Path) -> str:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _load_pdf(self, file_path: Path) -> str:
        """加载 PDF 文件"""
        reader = pypdf.PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        return "\n\n".join(text_parts)
    
    def _load_docx(self, file_path: Path) -> str:
        """加载 Word 文档 (.docx)"""
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_doc(self, file_path: Path) -> str:
        """加载旧版 Word 文档 (.doc)"""
        # 方法1: 尝试使用 antiword (Linux系统)
        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
            pass

        # 方法2: 尝试使用 catdoc (某些系统)
        try:
            result = subprocess.run(
                ['catdoc', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
            pass

        # 方法3: 尝试使用 LibreOffice 转换为 docx 再读取
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                # 使用 LibreOffice 转换
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'docx',
                     '--outdir', tmpdir, str(file_path)],
                    capture_output=True,
                    timeout=60
                )
                if result.returncode == 0:
                    # 找到生成的 docx 文件
                    docx_path = Path(tmpdir) / (file_path.stem + '.docx')
                    if docx_path.exists():
                        return self._load_docx(docx_path)
        except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
            pass

        # 如果所有方法都失败
        raise ValueError(
            "无法读取 .doc 文件。请安装以下工具之一：\n"
            "- antiword: apt-get install antiword\n"
            "- catdoc: apt-get install catdoc\n"
            "- LibreOffice: apt-get install libreoffice"
        )
    
    def _load_xlsx(self, file_path: Path) -> str:
        """加载 Excel 文件"""
        # 读取所有工作表
        excel_file = pd.ExcelFile(file_path)
        all_text = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            # 将表格转换为文本格式
            text = f"【工作表: {sheet_name}】\n"
            text += df.to_string(index=False, na_rep='')
            all_text.append(text)
        
        return "\n\n".join(all_text)


def process_file(file_path: str, user_id: Optional[int] = None) -> Optional[Document]:
    """处理单个文件
    
    Args:
        file_path: 文件路径
        user_id: 用户ID（可选，用于标识）
    
    Returns:
        Document 对象或 None（处理失败时）
    """
    path = Path(file_path)
    if not path.exists():
        print(f"文件不存在: {file_path}")
        return None
    
    processor = DocumentProcessor(str(path.parent))
    
    try:
        docs = processor.load_document(path)
        if docs:
            doc = docs[0]
            # 添加用户ID到metadata（如果提供）
            if user_id is not None:
                doc.metadata['user_id'] = user_id
            return doc
    except Exception as e:
        print(f"处理文件失败 {path.name}: {e}")
    
    return None


def process_directory(input_dir: str) -> List[Document]:
    """处理目录中的所有文档"""
    processor = DocumentProcessor(input_dir)
    return processor.load_all_documents()


if __name__ == "__main__":
    # 测试
    docs = process_directory("./docs")
    print(f"共加载 {len(docs)} 个文档")
    for doc in docs[:3]:
        print(f"- {doc.metadata['filename']}: {len(doc.page_content)} 字符")